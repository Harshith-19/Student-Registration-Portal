import json
import pickle
import docx
import smtplib
from datetime import datetime
from email.message import EmailMessage
from flask import Flask, render_template, request, jsonify   
app = Flask(__name__)

@app.route("/")
def home():
    return render_template("index.html")    
    

@app.route("/register/", methods=["POST",'GET'])
def home1():
    return render_template("register.html") 
@app.route("/submitJSON1", methods=["POST",'GET'])
def processJSON1():
    err=''
    flag=0
    jsonStr = request.get_json()
    jsonObj = json.loads(jsonStr) 
    
    response = ""
    name=jsonObj['name']
    gmail=jsonObj['gmail']
    if '.' not in gmail and '@' not in gmail:
        err+="<b>Please enter a valid email</b><br><br>"
        flag=1
    fname=jsonObj['fname']
    mname=jsonObj['mname']
    gender=jsonObj['gender']
    dob=jsonObj['dob']
    category=jsonObj['category']
    religion=jsonObj['religion']
    mob=jsonObj['mob']
    try:
        xyz=int(mob)
    except:
        flag=1
        err+="<b>Please enter a valid mobile no</b><br><br>"
    hno=jsonObj['hno']
    cv=jsonObj['cv']
    district=jsonObj['district']
    state=jsonObj['state']
    pincode=jsonObj['pincode']
    try:
        xyz=int(pincode)
    except:
        flag=1
        err+="<b>Please enter a valid pincode</b><br><br>"
    country=jsonObj['country']
    age=jsonObj['age']
    try:
        xyz=int(age)
    except:
        flag=1
        err+="<b>Please enter a valid age</b><br><br>"
    

    msg=EmailMessage()
    msg['Subject']='Exam Registration'
    msg['From']='studentic574@gmail.com'
    msg['To']=gmail
    f=open("b.txt","w")
    f.write(f"Dear {name} \n")
    f.close()
    f=open("b.txt","a")
    f.write("    You have successfully registered  for the exam.\nWe are attaching your registration form here.\nkeep studying.\n\n\nExam team\nGroup 35 ,Raipur\nMob No-7725932105\n ")
    f.close()
    #determining the total no of entries in the db file
    with open('count.txt','r') as f:
        no=int(f.read())+1
    
    with open("b.txt") as myfile:
        body=myfile.read()
        msg.set_content(body)
    myfile.close()
    b=datetime.today().strftime("%I:%M %p")
    a=str(datetime.today())
    today_date=a[:-7]+" "+b[-2:]


    doc = docx.Document()
    
    doc.add_heading('             Exam Registration Application', 0)
    doc.add_heading('                                                          Indian Institute of Technology Bhilai\n',3 )

    # table3=doc.add_table(rows=1,cols=2)
    # row=table3.rows[0].cells
    # row[0].text="Application Details"
    a=open("p.pickle","rb")
    l=pickle.load(a)
    Application_no=l[0]
    c=l[1]
    l[0]+=111111
    l[1]+=1
    a.close()

    f=open("p.pickle","wb")
    pickle.dump(l,f)
    f.close()
    data3=(("Application no",Application_no),("Application Submitted Date",today_date),("Exam","Tierce Exam"))
    dataapp=(("Application no",Application_no),("Application Submitted Date",today_date))
    # for j,k in data3:
    #     ror=table3.add_row().cells
    #     row[0].text=str(j)
    #     row[1].text=str(k)
    table3 = doc.add_table(rows=1, cols=2)
        
    row = table3.rows[0].cells
    row[0].text = 'Application Details'

    for i,j in data3:

        
    # for id, name in data3:
        
        row = table3.add_row().cells
        row[0].text = i
        row[1].text = str(j)
            
        
    table3.style = 'Colorful List'
    # table3.style="Colorful List"
    # doc.add_paragraph(' Indian Institute of Technology Bhilai' )
    doc.add_paragraph(' ' )


        

        
    data1 = (("Student Name", name),("Father's Name", fname), ("Mother's Name", mname),("Gender",gender),("Date of Birth ",dob),("Age",age) ,("Category",category),("Religion",religion),("Mobile No",mob),("Gmail id",gmail))
        
    table1 = doc.add_table(rows=1, cols=2)
        
    row = table1.rows[0].cells    
    row[0].text = 'Canditate Details'
        # row[1].text = ''
    for id, name in data1:
        
        row = table1.add_row().cells
        row[0].text = str(id)
        row[1].text = name
            
    table1.style = 'Colorful List'
    doc.add_paragraph(' ' )
    # doc.add_heading(" Candidate's  Address Deatils",4 )
    data2 = (("House no", hno),("City/Village", cv), ("District", district),("State",state),("Pincode ",pincode),("Country",country))
    #adding to the db file
    d1=dict([('Sno',no)]+list(dataapp)+list(data1)+list(data2))
    d12=dict(list(data1)+list(data2))
    with open('database.txt','rb') as d:
        for i in range(1,no):
            xyz1=pickle.load(d)
            xyz1.pop('Application no')
            xyz1.pop('Application Submitted Date')
            xyz1.pop('Sno')
            if xyz1==d12:
                flag=1
                err+="<b>Student already registered <b></b><br>"
                break
    if flag==0:
        with open('count.txt','w') as f:
            f.write(str(no))
        table2 = doc.add_table(rows=1, cols=2)
        with open('database.txt','ab') as d:
            pickle.dump(d1,d)
        row = table2.rows[0].cells
        row[0].text = "Candidate's  Address Deatils"
        
        for id, name in data2:
        
            row = table2.add_row().cells
            row[0].text = str(id)
            row[1].text = name
            
        
        table2.style = 'Colorful List'
        doc.add_paragraph(' ' )
        doc.add_heading("Declaration:",4 )
        doc.add_paragraph('    I hereby declare that all information given in this application form are true and correct to the best of my knowledge and belief. In the event of my information being found false or incorrect, action can be taken against me by the Institute.' )
        doc.save('registration_form.docx')
        with open('registration_form.docx','rb') as f:
            f1=f.read()
        f.close()
        msg.add_attachment(f1,maintype='application',subtype='docx',filename='registration_form.docx')
        server=smtplib.SMTP_SSL("smtp.gmail.com",465)
        server.login("studentic574@gmail.com","student@2002")
        server.send_message(msg)
        server.quit()
        
   
        response+="<b> You are now registered <b>""  ""</b><br>"
        response+="<b>""  ""</b><br>"
        response+="<b> Please check your email <b>""  ""</b><br>"
    	    
        return response
    else:
        response+=err
        return response
@app.route("/pyq/")
def home2():
    return render_template("pyq.html")
@app.route("/admin/")
def home3():
    return render_template("admin.html")
@app.route("/submitJSON2", methods=["POST",'GET'])
def processJSON2(): 
    jsonStr = request.get_json()
    jsonObj = json.loads(jsonStr) 
    
    response = ""
    id=jsonObj['id']
    password=jsonObj['pass']
    with open('password.txt','r') as d:
        pswd=d.read()
    if not (id=='12041580' and password==pswd):
        response+="<b> Wrong Credentials Please Try again <b>""  ""</b><br>"
        return response
    else :
        response+="<b>Database</b><br><br>"
        d=open("database.txt",'rb')
    while True:
        try:
            response+="<b> "+str(pickle.load(d))+"</b><br><br><br>"
        except:
            break
    d.close()
    return response
@app.route("/admitcard/")
def home4():
    return render_template("admitcard.html")
@app.route("/submitJSON3", methods=["POST",'GET'])
def processJSON3(): 
    jsonStr = request.get_json()
    jsonObj = json.loads(jsonStr) 
    
    response = ""
    appno=int(jsonObj['appno'])
    with open('count.txt','r') as f:
        z=int(f.read())
    f=open('database.txt','rb')
    g=""
    flag=0
    for i in range(0,z):
        k=pickle.load(f)
        if k["Application no"]==appno:
            g+="<b>found you</b><br><br>"
            g+="<b>"+str(k)+"</b><br><br>"


            doc = docx.Document()
  
            doc.add_heading('                             Indian Institute of Technology Bhilai', 1)
            doc.add_heading('                                                    Tierce Exam -Admit Card\n',3 )


            

            doc.add_paragraph(' ' )
            # doc.add_heading(" Candidate's  Address Deatils",4 )
            data4 = (("Application_no", str(k["Application no"])), ("Student Name", str(k["Student Name"])),("Father name",str(k["Father's Name"])),("Gender",str(k["Gender"])),("Date of Birth",str(k["Date of Birth "])),("Category",str(k["Category"])))

            table1 = doc.add_table(rows=1, cols=2)

            row = table1.rows[0].cells
            row[0].text = "                                             Student Details"

            for id, name in data4:
                
                row = table1.add_row().cells
                row[0].text = str(id)
                row[1].text = name

            table1.style = 'Colorful List'
            doc.add_paragraph(' ' )

            data5 = (("Question Paper Medium", "English"),("Course Name", "IC100"), ("Date of Examination", "14-09-2021"),("Shift","first"),("Reporting ","9:00 AM"),("Gate Closing","9:45 AM"),("Timing of Exam","10:00 AM"),("Exam Room","207"),("Address","Indian institute of Technology Bhilai,GEC Campus,Sejbahar,Raipur-492015 (C.G.)"))

            table2 = doc.add_table(rows=1, cols=2)
                
            row = table2.rows[0].cells
            row[0].text = "                                             Exam Details"
                
            for id, name in data5:
                
                row = table2.add_row().cells
                row[0].text = str(id)
                row[1].text = name
                
            table2.style = 'Colorful List'
            doc.add_paragraph(' ' )
            doc.add_heading("Undertaking:",4 )
            doc.add_paragraph('    That, I have read the instructions, Guidelines and relevant orders pertaining to COVID-19 pandemic. I have read the information Bulletin, Instructions and Notices related to this examination available on the website.' )

            doc.save('admitcard.docx')
            stname=k["Student Name"]

            msg=EmailMessage()
            msg['Subject']='Exam Registration'
            msg['From']='studentic574@gmail.com'
            msg['To']=k["Gmail id"]
            f=open("c.txt","w")
            f.write(f"Dear {stname}\n")
            f.close()
            f=open("c.txt","a")
            f.write("    We are attaching your admit card here.\nFollow the exam guidelines ,rule and regulations.    \nkeep studying.\n\n\nExam team\nGroup 35 ,Raipur\nMob No-7725932105\n ")
            f.close()

            with open("c.txt") as myfile:
                body=myfile.read()
                msg.set_content(body)
                myfile.close()
            
            with open('admitcard.docx','rb') as f:
                f2=f.read()
                
            msg.add_attachment(f2,maintype='application',subtype='docx',filename='admitcard.docx')
            server=smtplib.SMTP_SSL("smtp.gmail.com",465)
            server.login("studentic574@gmail.com","student@2002")
            server.send_message(msg)
            server.quit()
        

            flag=1
            break
    if flag==1:
        response+=g
        response+="<b>We have sent your admit card on your registered gmail id</b><br><br>"

    else:
        response+="<b>Application no is wrond</b><br><br>"
    return response
@app.route("/changepassword/")
def home5():
    return render_template("changepassword.html")
@app.route("/submitJSON5", methods=["POST",'GET'])
def processJSON5(): 
    jsonStr = request.get_json()
    jsonObj = json.loads(jsonStr) 
    response = ""
    p1=jsonObj['p1']
    p2=jsonObj['p2']
    p3=jsonObj['p3']
    err=''
    flag=0
    with open('password.txt','r') as d:
        pswd=d.read()
    if p1!=pswd:
        flag=1
        err+="<b>Worng password</b><br><br>"
    if p2!=p3:
        flag=1
        err+="<b>Two passwords are not matching</b><br><br>"
    if flag==1:
        response+=err
        return response
    else:
        with open('password.txt','w') as d:
            d.write(p2)
        response+="<b>Password Changed</b><br><br>"
        return response
        


            



 

    
    
if __name__ == "__main__":
    app.run(debug=True)
    
    
